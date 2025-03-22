import os
from PyPDF2 import PdfReader
import docx
from models import Document, DocumentMeta, Session
from rag import add_documents_to_collection, generate_summary

class DocumentProcessor:
    def __init__(self):
        self.session = Session()
    
    def process_file(self, file):
        """
        Process uploaded file
        
        Args:
            file: Uploaded file object
        """
        filename = file.filename
        title = os.path.splitext(filename)[0]
        content = self.extract_content(file)
        summary = generate_summary(content)
        
        # Save document content
        doc = Document(content=content)
        self.session.add(doc)
        self.session.flush()
        
        # Save metadata
        meta = DocumentMeta(
            id=doc.id,
            title=title,
            summary=summary
        )
        self.session.add(meta)
        self.session.commit()
        
        # Add to ChromaDB
        add_documents_to_collection(
            documents=[content],
            ids=[f"doc_{doc.id}"]
        )
        
        return doc.id
    
    def extract_content(self, file):
        ext = os.path.splitext(file.filename)[1].lower()
        if ext == '.pdf':
            return self._extract_pdf(file)
        elif ext == '.txt':
            return file.read().decode('utf-8')
        elif ext in ['.doc', '.docx']:
            return self._extract_docx(file)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_pdf(self, file):
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    
    def _extract_docx(self, file):
        doc = docx.Document(file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def delete_document(self, doc_id):
        """
        Delete document with specified ID
        
        Args:
            doc_id: Document ID
        """
        # Delete document from ChromaDB
        from rag import collection
        try:
            collection.delete(ids=[f"doc_{doc_id}"])
        except:
            pass  # Ignore error if document doesn't exist in ChromaDB
        
        self.session.query(Document).filter_by(id=doc_id).delete()
        self.session.query(DocumentMeta).filter_by(id=doc_id).delete()
        self.session.commit()
    
    def get_documents(self, page=1, per_page=10):
        """
        Get document list with pagination support
        
        Args:
            page: Page number, starting from 1
            per_page: Number of items per page
        
        Returns:
            tuple: (document list, total count)
        """
        offset = (page - 1) * per_page
        total = self.session.query(DocumentMeta).count()
        metas = self.session.query(DocumentMeta)\
            .order_by(DocumentMeta.id.desc())\
            .offset(offset)\
            .limit(per_page)\
            .all()
        return metas, total 